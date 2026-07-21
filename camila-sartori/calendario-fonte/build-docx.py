# -*- coding: utf-8 -*-
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x5A, 0x10, 0x42)
ORANGE = RGBColor(0x9C, 0x12, 0x56)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CARDGRAY = "FBEEF7"
OUTERBOX = "FDF6FB"
CH_HEX = {'INSTAGRAM': '9C1256', 'WHATSAPP': '1F8A4C'}

TIPO_EMOJI = [
    ('REEL', '🎬'), ('CARROSSEL', '🎯'), ('VENDA', '💰'),
    ('ENGAJAMENTO', '💬'), ('EXCLUSIVO', '✨'), ('REPOST', '🔁'), ('ENQUETE', '📊'),
]

def tipo_emoji(tipo):
    for key, emoji in TIPO_EMOJI:
        if key.upper() in tipo.upper():
            return emoji
    return '📌'

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(section, m, Cm(1.5))

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def shade_paragraph(p, hexcolor, border=True):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    pPr.append(shd)
    if border:
        pBdr = OxmlElement('w:pBdr')
        for edge in ('top', 'left', 'bottom', 'right'):
            el = OxmlElement('w:%s' % edge)
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '6')
            el.set(qn('w:space'), '4')
            el.set(qn('w:color'), 'D9D9D9')
            pBdr.append(el)
        pPr.append(pBdr)

def add_run(p, text, bold=False, color=None, size=10, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r

def add_break(p):
    p.add_run().add_break()

def month_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(14)
    add_run(p, text, bold=True, color=NAVY, size=15)

def week_bar(text):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.columns[0].width = Cm(18)
    cell = t.rows[0].cells[0]
    cell.width = Cm(18)
    shade_cell(cell, 'E09C3B')
    p = cell.paragraphs[0]
    add_run(p, text, bold=True, color=WHITE, size=11.5)
    sp_after = doc.add_paragraph()
    sp_after.paragraph_format.space_after = Pt(4)

def day_bar(text):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.columns[0].width = Cm(18)
    cell = t.rows[0].cells[0]
    cell.width = Cm(18)
    shade_cell(cell, '0A0C0F')
    p = cell.paragraphs[0]
    add_run(p, "📅  " + text, bold=True, color=WHITE, size=10.5)
    sp_after = doc.add_paragraph()
    sp_after.paragraph_format.space_after = Pt(4)

def set_cell_borders(cell, color='D9D9D9', sz='8'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:%s' % edge)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)

def begin_post_box(time, channel, tipo):
    """Cria UMA unica tabela de 2 linhas (cabecalho colorido + corpo) para o post inteiro,
    com borda ao redor de tudo, igual ao modelo de junho: uma caixa so, sem espaco entre o
    cabecalho e o corpo."""
    hexcolor = CH_HEX.get(channel, '0A0C0F')
    emoji = tipo_emoji(tipo)
    t = doc.add_table(rows=2, cols=1)
    t.autofit = False
    t.columns[0].width = Cm(18)
    head_cell = t.rows[0].cells[0]
    head_cell.width = Cm(18)
    shade_cell(head_cell, hexcolor)
    set_cell_borders(head_cell)
    p = head_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "⏰ %s   [ %s ]   %s %s" % (time, channel, emoji, tipo), bold=True, color=WHITE, size=10)

    body_cell = t.rows[1].cells[0]
    body_cell.width = Cm(18)
    shade_cell(body_cell, OUTERBOX)
    set_cell_borders(body_cell)
    body_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    return body_cell

def card_paragraph(cell, label, lines, label_color=ORANGE):
    """lines: list of (bold_prefix_or_None, text). Renders as ONE shaded/bordered paragraph
    (label + lines joined by line breaks), so it reads as a single visual card."""
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    shade_paragraph(p, CARDGRAY)
    add_run(p, label, bold=True, color=label_color, size=9)
    for prefix, text in lines:
        add_break(p)
        if prefix:
            add_run(p, prefix + " ", bold=True, size=9.5)
        if text:
            add_run(p, text, size=9.5)

def capa_lines(headline, subheadline, cta):
    lines = [("HEADLINE:", headline)]
    if subheadline:
        lines.append(("SUBHEADLINE:", subheadline))
    if cta:
        lines.append(("CTA:", cta))
    return lines

def field_p(cell, label, text, color=None, size=10, italic=False, emoji=None):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    prefix = (emoji + "  ") if emoji else ""
    add_run(p, prefix + label + " ", bold=True, size=size, color=color, italic=italic)
    add_run(p, text, size=size, color=color, italic=italic)

def post_card(time, channel, tipo, formato, headline, subheadline, cta,
              legenda=None, ancoragem=None, persona=None, hashtags=None, origem=None,
              cards=None, card_final=None, observacao=None, cenario=None,
              label_legenda="LEGENDA:", acoes=None):
    cell = begin_post_box(time, channel, tipo)

    if origem:
        field_p(cell, "Origem:", origem, color=GRAY, size=8.5, italic=True, emoji="📖")
    if ancoragem:
        field_p(cell, "Ancoragem:", ancoragem, color=GRAY, size=8.5, italic=True, emoji="📚")
    if cenario:
        field_p(cell, "Cenário sugerido:", cenario, color=GRAY, size=8.5, italic=True, emoji="🎬")
    formato_persona = formato + ("     Persona: " + persona if persona else "")
    field_p(cell, "Formato:", formato_persona, size=9)

    if cards:
        card_paragraph(cell, "CARD 1 — CAPA", capa_lines(headline, subheadline, cta))
        for i, c in enumerate(cards, start=2):
            lines = [(c['titulo'], '')] if c.get('titulo') else []
            lines.append((None, c['corpo']))
            card_paragraph(cell, "CARD %d" % i, lines, label_color=GRAY)
        if card_final:
            card_paragraph(cell, "CARD FINAL", capa_lines(card_final.get('headline'), card_final.get('subheadline'), card_final.get('cta')))
        field_p(cell, "LEGENDA DO POST:", legenda, size=9.5)
    else:
        if headline:
            field_p(cell, "HEADLINE:", headline, size=10.5)
        if subheadline:
            field_p(cell, "SUBHEADLINE:", subheadline, size=9.5)
        if cta:
            field_p(cell, "CTA:", cta, color=ORANGE, size=9.5)
        if acoes:
            for i, a in enumerate(acoes, start=1):
                card_paragraph(cell, "AÇÃO %d — %s" % (i, a['label']), [(None, a['texto'])])
        elif legenda:
            field_p(cell, label_legenda, legenda, size=9.5)

    if hashtags:
        field_p(cell, "HASHTAGS:", hashtags, color=GRAY, size=8.5)
    if observacao:
        field_p(cell, "Observação:", observacao, color=ORANGE, size=8.5, italic=True, emoji="📌")

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(12)

def podcast_entry(time, episodio, ultimo=False):
    cell = begin_post_box(time, "INSTAGRAM", "ElevaCast — Corte de Podcast (Reels)")
    texto = ("Reels — Corte do ElevaCast · Episódio: %s — %s. "
             "Publicação programada, não requer aprovação de copy." %
             (episodio, "Último corte da série" if ultimo else "Corte sequencial"))
    field_p(cell, "Formato:", texto, size=9.5)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(12)

def add_month(title):
    month_title(title)

def add_week(title):
    week_bar(title)

def add_day(title, posts):
    day_bar(title)
    for p in posts:
        post_card(**p)

def add_podcast(title, time, episodio, ultimo=False):
    day_bar(title)
    podcast_entry(time, episodio, ultimo=ultimo)

content_file = sys.argv[1]
output_docx = sys.argv[2]
namespace = dict(add_month=add_month, add_week=add_week, add_day=add_day, add_podcast=add_podcast)
exec(open(content_file, encoding='utf-8').read(), namespace)

doc.save(output_docx)
print("DOCX gerado:", output_docx)
